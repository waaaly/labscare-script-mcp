"""
DOCX to XMReport Converter
将docx文档转换为labscare自定义的xmreport格式
支持doc格式自动转换为docx后再转换
"""

import zipfile
import os
import sys
import base64
import re
import uuid
import subprocess
import tempfile
import shutil
from xml.etree import ElementTree as ET
from xml.dom import minidom
from typing import Optional
from dataclasses import dataclass, field
from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

def _get_subprocess_kwargs():
    """Windows 下完全隐藏所有子进程窗口"""
    kwargs = {}
    if sys.platform == 'win32':
        si = subprocess.STARTUPINFO()
        si.dwFlags |= subprocess.STARTF_USESHOWWINDOW
        si.wShowWindow = subprocess.SW_HIDE  # 隐藏窗口，而非 CREATE_NO_WINDOW
        kwargs['startupinfo'] = si
        kwargs['creationflags'] = (
            subprocess.CREATE_NO_WINDOW |
            subprocess.DETACHED_PROCESS  # 与父进程分离，防止继承控制台
        )
    return kwargs

def convert_doc_to_docx(doc_path: str, output_dir: str = None) -> str:
    """
    将doc文件转换为docx格式
    
    Args:
        doc_path: doc文件路径
        output_dir: 输出目录（可选，默认为doc文件所在目录）
    
    Returns:
        转换后的docx文件路径
    
    Raises:
        RuntimeError: 如果转换失败
    """
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"文件不存在: {doc_path}")
    
    if output_dir is None:
        output_dir = os.path.dirname(doc_path) or '.'
    
    docx_path = os.path.join(output_dir, os.path.splitext(os.path.basename(doc_path))[0] + '.docx')
    
    soffice_paths = [
        'soffice',
        '/usr/bin/soffice',
        '/usr/lib/libreoffice/program/soffice',
        'C:\\Program Files\\LibreOffice\\program\\soffice.exe',
        'C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe',
    ]
    
    soffice_cmd = None
    for path in soffice_paths:
        try:
            # Windows下隐藏控制台窗口
            kwargs = {}
            if sys.platform == 'win32':
                kwargs['creationflags'] = subprocess.CREATE_NO_WINDOW
            subprocess.run(
                [path, '--version'],
                capture_output=True,
                stdin=subprocess.DEVNULL,
                timeout=5,
                **_get_subprocess_kwargs()   # ← 替换原来的 kwargs
            )
            soffice_cmd = path
            break
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
    
    if soffice_cmd:
        print(f"使用LibreOffice转换: {doc_path}")
        try:
            # Windows下隐藏控制台窗口
            kwargs = {}
            if sys.platform == 'win32':
                kwargs['creationflags'] = subprocess.CREATE_NO_WINDOW
            
           # 执行转换
                result = subprocess.run(
                    [
                        soffice_cmd,
                        '--headless',
                        '--norestore',
                        '--nofirststartwizard',
                        '--nolockcheck',
                        '--convert-to', 'docx',
                        '--outdir', output_dir,
                        doc_path
                    ],
                    capture_output=True,
                    text=True,
                    timeout=120,
                    stdin=subprocess.DEVNULL,
                    **_get_subprocess_kwargs()   # ← 替换原来的 kwargs
                )
            if result.returncode == 0 and os.path.exists(docx_path):
                print(f"LibreOffice转换成功: {docx_path}")
                return docx_path
            else:
                print(f"LibreOffice转换失败: {result.stderr}")
        except subprocess.TimeoutExpired:
            print("LibreOffice转换超时")
        except Exception as e:
            print(f"LibreOffice转换出错: {e}")
    
    try:
        import win32com.client
        print(f"使用Microsoft Word转换: {doc_path}")

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        try:
            doc = word.Documents.Open(os.path.abspath(doc_path))
            wdFormatXMLDocument = 12
            doc.SaveAs2(os.path.abspath(docx_path), FileFormat=wdFormatXMLDocument)
            doc.Close(False)
            if os.path.exists(docx_path) and zipfile.is_zipfile(docx_path):
                print(f"Microsoft Word转换成功: {docx_path}")
                return docx_path
            raise RuntimeError(f"Microsoft Word输出不是有效docx文件: {docx_path}")
        finally:
            word.Quit()
    except ImportError:
        print("未安装pywin32，无法使用Microsoft Word转换")
    except Exception as e:
        print(f"Microsoft Word转换出错: {e}")
    
    raise RuntimeError(
        "doc转docx失败。请安装以下任一工具：\n"
        "1. LibreOffice (推荐，跨平台)\n"
        "2. Microsoft Word + pywin32 (仅Windows)\n"
        "   安装pywin32: pip install pywin32"
    )


@dataclass
class TextContent:
    text: str = ""
    x: int = 0
    y: int = 0
    width: int = 0
    height: int = 0


@dataclass
class ImageContent:
    image_data: str = ""
    x: int = 0
    y: int = 0
    width: int = 0
    height: int = 0


@dataclass
class TableCell:
    content: str = ""
    col_span: int = 1
    row_span: int = 1
    images: list = field(default_factory=list)
    children: list = field(default_factory=list)


@dataclass
class TableRow:
    cells: list = field(default_factory=list)
    height: float = 0


@dataclass
class TableContent:
    rows: list = field(default_factory=list)
    x: int = 0
    y: int = 0
    width: int = 0
    column_widths: list = field(default_factory=list)


@dataclass
class PageContent:
    header: list = field(default_factory=list)
    body: list = field(default_factory=list)
    footer: list = field(default_factory=list)
    header_images: list = field(default_factory=list)


class DocxParser:
    NS = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    }

    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.images = {}
        self._extract_images()

    def _extract_images(self):
        with zipfile.ZipFile(self.docx_path, 'r') as z:
            for name in z.namelist():
                if name.startswith('word/media/') and not name.endswith('/'):
                    image_data = z.read(name)
                    image_ext = os.path.splitext(name)[1].lower()
                    mime_type = self._get_mime_type(image_ext)
                    if mime_type:
                        b64_data = base64.b64encode(image_data).decode('utf-8')
                        self.images[name] = f"data:{mime_type};base64,{b64_data}"

    def _extract_header_images(self) -> list:
        header_images = []
        with zipfile.ZipFile(self.docx_path, 'r') as z:
            header_files = sorted([
                n for n in z.namelist()
                if n.startswith('word/header') and n.endswith('.xml')
            ])
            for header_xml in header_files:
                rels_xml = f"word/_rels/{os.path.basename(header_xml)}.rels"
                rel_map = {}
                if rels_xml in z.namelist():
                    rel_root = ET.fromstring(z.read(rels_xml))
                    for rel in rel_root.findall('{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                        rid = rel.get('Id')
                        target = rel.get('Target')
                        if rid and target:
                            rel_map[rid] = target

                header_root = ET.fromstring(z.read(header_xml))
                for imagedata in header_root.findall('.//{urn:schemas-microsoft-com:vml}imagedata'):
                    rid = imagedata.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    target = rel_map.get(rid)
                    if not target:
                        continue
                    media_path = f"word/{target.lstrip('./').replace('..\\\\', '').replace('../', '')}"
                    media_path = media_path.replace('word/word/', 'word/')
                    img_data = self.images.get(media_path)
                    if img_data and img_data not in header_images:
                        header_images.append(img_data)
        return header_images

    def _get_mime_type(self, ext: str) -> Optional[str]:
        mime_map = {
            '.png': 'image/png',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.gif': 'image/gif',
            '.bmp': 'image/bmp',
            '.tiff': 'image/tiff',
            '.emf': 'image/x-emf',
            '.wmf': 'image/x-wmf',
        }
        return mime_map.get(ext)

    def _get_paragraph_text(self, para) -> str:
        return para.text.strip()

    def _extract_cell_images(self, cell) -> list:
        images = []
        tc = cell._tc
        blips = tc.xpath('.//a:blip')
        for blip in blips:
            embed = blip.get(qn('r:embed'))
            if not embed:
                continue
            rel = cell.part.rels.get(embed)
            if not rel:
                continue
            partname = str(rel.target_part.partname).lstrip('/')
            img_data = self.images.get(partname)
            if img_data:
                images.append(img_data)
        return images

    def _get_table_column_widths(self, table) -> list:
        widths = []
        tbl = table._tbl
        tbl_grid = tbl.tblGrid
        if tbl_grid is not None and tbl_grid.gridCol_lst:
            for gc in tbl_grid.gridCol_lst:
                w = gc.w
                if w is not None:
                    widths.append(int(w))

        if not widths:
            for col in table.columns:
                try:
                    if col.width:
                        widths.append(int(col.width.twips))
                except Exception:
                    pass

        return widths

    def _get_table_data(self, table) -> TableContent:
        result = TableContent()
        result.column_widths = self._get_table_column_widths(table)

        vmerge_track = {}

        for row in table.rows:
            table_row = TableRow()

            for col_idx, cell in enumerate(row.cells):
                cell_content = ""
                for para in cell.paragraphs:
                    para_text = para.text.strip()
                    if para_text:
                        cell_content = f"{cell_content}\n{para_text}".strip() if cell_content else para_text

                col_span = 1
                row_span = 1

                tc = cell._tc
                tcPr = tc.tcPr
                if tcPr is not None:
                    gridSpan = tcPr.find(qn('w:gridSpan'))
                    if gridSpan is not None:
                        col_span = int(gridSpan.get(qn('w:val'), 1))

                    vMerge = tcPr.find(qn('w:vMerge'))
                    if vMerge is not None:
                        merge_val = vMerge.get(qn('w:val')) or 'continue'
                        if merge_val == 'restart':
                            vmerge_track[col_idx] = 1
                            row_span = 1
                        else:
                            if col_idx in vmerge_track:
                                vmerge_track[col_idx] += 1
                            row_span = 1

                table_cell = TableCell(
                    content=cell_content,
                    col_span=col_span,
                    row_span=row_span,
                    images=self._extract_cell_images(cell)
                )
                table_row.cells.append(table_cell)

            result.rows.append(table_row)

        return result

    def parse(self) -> PageContent:
        page = PageContent()
        page.header_images = self._extract_header_images()
        
        section = self.doc.sections[0] if self.doc.sections else None
        
        header_content = []
        # 使用xmreport固定页眉骨架，避免docx原始header结构干扰
        page.header = header_content
        
        footer_content = []
        if section and section.footer:
            for para in section.footer.paragraphs:
                text = self._get_paragraph_text(para)
                if text:
                    footer_content.append(TextContent(text=text))
            for table in section.footer.tables:
                footer_content.append(self._get_table_data(table))
        page.footer = footer_content
        
        body_content = []
        for element in self.doc.element.body:
            if element.tag == qn('w:p'):
                for para in self.doc.paragraphs:
                    if para._element == element:
                        text = self._get_paragraph_text(para)
                        if text:
                            body_content.append(TextContent(text=text))
                        break
            elif element.tag == qn('w:tbl'):
                for table in self.doc.tables:
                    if table._element == element:
                        body_content.append(self._get_table_data(table))
                        break
        page.body = body_content
        
        return page

    def get_page_settings(self) -> dict:
        section = self.doc.sections[0] if self.doc.sections else None
        if not section:
            return {
                'width': 2090,
                'height': 2960,
                'left_margin': 154,
                'right_margin': 154,
                'top_margin': 20,
                'bottom_margin': 118,
                'orientation': 'portait'
            }
        
        width = int(section.page_width.pt * 10) if section.page_width else 2090
        height = int(section.page_height.pt * 10) if section.page_height else 2960
        
        left_margin = int(section.left_margin.pt * 10) if section.left_margin else 154
        right_margin = int(section.right_margin.pt * 10) if section.right_margin else 154
        top_margin = int(section.top_margin.pt * 10) if section.top_margin else 20
        bottom_margin = int(section.bottom_margin.pt * 10) if section.bottom_margin else 118
        
        orientation = 'portait'
        if section.orientation == WD_ORIENT.LANDSCAPE:
            orientation = 'landscape'
        
        return {
            'width': width,
            'height': height,
            'left_margin': left_margin,
            'right_margin': right_margin,
            'top_margin': top_margin,
            'bottom_margin': bottom_margin,
            'orientation': orientation
        }


class XmreportGenerator:
    def __init__(self):
        self.id_counter = 0
        self.ids = {}

    def _generate_id(self, prefix: str = "") -> str:
        self.id_counter += 1
        # 生成与系统导出格式一致的ID
        # 格式: 19C + 10位随机字符 + - + 8位随机字符
        unique1 = uuid.uuid4().hex[:10].upper()
        unique2 = uuid.uuid4().hex[:8].upper()
        return f"19C{unique1}-{unique2}"

    def _create_element(self, tag: str, text: str = "") -> ET.Element:
        elem = ET.Element(tag)
        if text:
            elem.text = text
        return elem

    def _create_data_bind(self, bind_text: str = "") -> ET.Element:
        data_bind = self._create_element("data-bind")
        data_bind.set("_id", self._generate_id("DB"))
        
        bind = self._create_element("bind")
        bind.text = bind_text
        data_bind.append(bind)
        
        for script_name in ["prepare-script", "post-script", "final-script"]:
            data_bind.append(self._create_element(script_name))
        
        return data_bind

    def _create_border(self) -> ET.Element:
        border = self._create_element("border")
        border.set("_id", self._generate_id("BD"))
        for side in ["left", "right", "top", "bottom", "tr2bl", "tl2br"]:
            border.append(self._create_element(side))
        return border

    def _create_background(self, transparent: bool = True) -> ET.Element:
        background = self._create_element("background")
        background.set("_id", self._generate_id("BG"))
        bg_type = self._create_element("type")
        bg_type.text = "TRANSPARENT" if transparent else "SOLID"
        background.append(bg_type)
        transparent_elem = self._create_element("transparent")
        transparent_elem.text = str(transparent).lower()
        background.append(transparent_elem)
        return background

    def _create_font(self) -> ET.Element:
        font = self._create_element("font")
        font.set("_id", self._generate_id("FT"))
        
        font.append(self._create_element("chinese-font", "宋体"))
        font.append(self._create_element("western-font", "Times New Roman"))
        
        for attr in ["bold", "italic", "under-line", "delete-line", "over-line"]:
            font.append(self._create_element(attr, "false"))
        
        for style in ["under-line-style", "delete-line-style", "over-line-style"]:
            font.append(self._create_element(style))
        
        font_size = self._create_element("font-size")
        alias = self._create_element("alias", "11pt")
        pounds = self._create_element("pounds", "11")
        font_size.append(alias)
        font_size.append(pounds)
        font.append(font_size)
        
        color = self._create_element("color", "rgba(0,0,0,1)")
        font.append(color)
        
        font.append(self._create_element("highlight", "false"))
        font.append(self._create_element("highlight-color", "rgba(255,255,255,1)"))
        
        return font

    def _create_paragraph(self) -> ET.Element:
        para = self._create_element("paragraph")
        para.set("_id", self._generate_id("PG"))
        
        para.append(self._create_element("writing-mode", "HORIZONTAL_TB"))
        para.append(self._create_element("text-align", "CENTER"))
        para.append(self._create_element("vertical-align", "MIDDLE"))
        para.append(self._create_element("text-indent", "0"))
        para.append(self._create_element("hanging-indent", "0"))
        para.append(self._create_element("line-height", "1"))
        para.append(self._create_element("line-height-unit", "LINE"))
        para.append(self._create_element("space-before", "0"))
        para.append(self._create_element("space-before-unit", "LINE"))
        para.append(self._create_element("space-after", "0"))
        para.append(self._create_element("space-after-unit", "LINE"))
        para.append(self._create_element("allow-breaking", "false"))
        para.append(self._create_element("shrink", "false"))
        para.append(self._create_element("letter-spacing", "0"))
        para.append(self._create_element("word-spacing", "0"))
        para.append(self._create_element("auto-space-dn", "false"))
        para.append(self._create_element("auto-space-de", "false"))
        para.append(self._create_element("line-wrap", "true"))
        para.append(self._create_element("cells-filling", "false"))
        para.append(self._create_element("cells", "0"))
        
        cell_line_style = self._create_element("cell-line-style")
        line_type = self._create_element("line-type")
        for part in ["aside", "middle", "bside"]:
            part_elem = self._create_element(part)
            if part == "middle":
                dash = self._create_element("dash")
                width = self._create_element("width", "1.7638888888888888")
                dummy = self._create_element("dummy", "false")
                part_elem.append(dash)
                part_elem.append(width)
                part_elem.append(dummy)
            line_type.append(part_elem)
        cell_line_style.append(line_type)
        cell_line_style.append(self._create_element("color", "rgba(0,0,0,1)"))
        para.append(cell_line_style)
        
        return para

    def _create_format(self) -> ET.Element:
        format_elem = self._create_element("format")
        format_elem.set("_id", self._generate_id("FM"))
        format_elem.append(self._create_element("format-type", "TEXT"))
        format_elem.append(self._create_element("format-category", "GENERAL"))
        format_elem.append(self._create_element("format-expression", "@"))
        return format_elem

    def _create_padding(self) -> ET.Element:
        padding = self._create_element("padding")
        padding.set("_id", self._generate_id("PD"))
        padding.append(self._create_element("left", "10"))
        padding.append(self._create_element("right", "10"))
        padding.append(self._create_element("top", "0"))
        padding.append(self._create_element("bottom", "0"))
        return padding

    def _create_text_element(self, text_content: TextContent, data_bind_id: str, 
                             border_id: str, bg_id: str, font_id: str, 
                             para_id: str, format_id: str, padding_id: str) -> ET.Element:
        text_elem = self._create_element("text")
        text_elem.append(self._create_element("id", f"text_{self.id_counter}"))
        text_elem.append(self._create_element("parent-id", "Body"))
        text_elem.append(self._create_element("z-index", "100"))
        text_elem.append(self._create_element("visible", "true"))
        text_elem.append(self._create_element("rotate", "0"))
        text_elem.append(self._create_element("adjust-height", "true"))
        text_elem.append(self._create_element("appear-in-every-page", "false"))
        text_elem.append(self._create_element("cross-page", "false"))
        
        bound = self._create_element("bound")
        bound.append(self._create_element("x", str(text_content.x)))
        bound.append(self._create_element("y", str(text_content.y)))
        bound.append(self._create_element("width", str(text_content.width or 1782)))
        bound.append(self._create_element("height", str(text_content.height or 30)))
        text_elem.append(bound)
        
        text_data_bind = self._create_element("data-bind")
        text_data_bind.set("_id", self._generate_id("TDB"))
        text_bind = self._create_element("bind")
        text_bind.text = text_content.text
        text_data_bind.append(text_bind)
        text_data_bind.append(self._create_element("prepare-script"))
        text_data_bind.append(self._create_element("post-script"))
        text_data_bind.append(self._create_element("final-script"))
        text_elem.append(text_data_bind)
        
        border_ref = self._create_element("border")
        border_ref.set("_ref", border_id)
        text_elem.append(border_ref)
        
        bg_ref = self._create_element("background")
        bg_ref.set("_ref", bg_id)
        text_elem.append(bg_ref)
        
        padding_ref = self._create_element("padding")
        padding_ref.set("_ref", padding_id)
        text_elem.append(padding_ref)
        
        font_ref = self._create_element("font")
        font_ref.set("_ref", font_id)
        text_elem.append(font_ref)
        
        para_ref = self._create_element("paragraph")
        para_ref.set("_ref", para_id)
        text_elem.append(para_ref)
        
        format_ref = self._create_element("format")
        format_ref.set("_ref", format_id)
        text_elem.append(format_ref)
        
        text_elem.append(self._create_element("text-type", "TEXT"))
        text_elem.append(self._create_element("use-rich-text-editor", "false"))
        
        return text_elem

    def _create_table_element(self, table_content: TableContent, data_bind_id: str,
                              border_id: str, bg_id: str, font_id: str,
                              para_id: str, format_id: str, padding_id: str) -> ET.Element:
        grid = self._create_element("grid")
        grid_id = f"grid{self.id_counter}"
        grid.append(self._create_element("id", grid_id))
        grid.append(self._create_element("parent-id", "Body"))
        grid.append(self._create_element("z-index", "100"))
        grid.append(self._create_element("visible", "true"))
        grid.append(self._create_element("appear-in-every-page", "false"))
        
        bound = self._create_element("bound")
        bound.append(self._create_element("x", str(table_content.x)))
        bound.append(self._create_element("y", str(table_content.y)))
        bound.append(self._create_element("width", str(table_content.width or 1782)))
        bound.append(self._create_element("height", "0"))
        grid.append(bound)
        
        data_bind_ref = self._create_element("data-bind")
        data_bind_ref.set("_ref", data_bind_id)
        grid.append(data_bind_ref)
        
        bg_ref = self._create_element("background")
        bg_ref.set("_ref", bg_id)
        grid.append(bg_ref)
        
        if table_content.column_widths:
            raw_widths = [float(w) for w in table_content.column_widths if w]
            if raw_widths:
                total_width = float(table_content.width or 1782)
                sum_raw = sum(raw_widths)
                scaled = [total_width * (w / sum_raw) for w in raw_widths]
                if scaled:
                    scaled[-1] = total_width - sum(scaled[:-1])
                col_widths_elem = self._create_element("column-widths")
                col_widths_elem.text = ",".join([str(round(w, 2)).rstrip('0').rstrip('.') for w in scaled])
                grid.append(col_widths_elem)
        
        grid.append(self._create_element("row-colors"))
        grid.append(self._create_element("col-colors"))
        
        rows_elem = self._create_element("rows")
        
        for row_idx, row in enumerate(table_content.rows):
            row_elem = self._create_element("row")
            row_id = f"row{self.id_counter + row_idx + 1}"
            row_elem.append(self._create_element("id", row_id))
            row_elem.append(self._create_element("parent-id", grid_id))
            row_elem.append(self._create_element("z-index", "100"))
            row_elem.append(self._create_element("visible", "true"))
            row_elem.append(self._create_element("rotate", "0"))
            row_elem.append(self._create_element("adjust-height", "true"))
            row_elem.append(self._create_element("appear-in-every-page", "false"))
            
            row_bound = self._create_element("bound")
            row_bound.append(self._create_element("x", "0"))
            row_bound.append(self._create_element("y", "0"))
            row_bound.append(self._create_element("width", "0"))
            row_bound.append(self._create_element("height", "0"))
            row_elem.append(row_bound)
            
            row_data_bind = self._create_element("data-bind")
            row_data_bind.set("_ref", data_bind_id)
            row_elem.append(row_data_bind)
            
            row_border = self._create_element("border")
            row_border.set("_ref", border_id)
            row_elem.append(row_border)
            
            row_bg = self._create_element("background")
            row_bg.set("_ref", bg_id)
            row_elem.append(row_bg)
            
            row_elem.append(self._create_element("enable-min-height", "false"))
            row_elem.append(self._create_element("min-height"))
            # 根据内容长度动态调整行高
            row_height = row.height if row.height > 0 else 33.33
            row_elem.append(self._create_element("height", str(row_height)))
            row_elem.append(self._create_element("adjust-height", "true"))
            row_elem.append(self._create_element("appear-in-every-page", "false"))
            row_elem.append(self._create_element("row-type", "SIMPLE"))
            row_elem.append(self._create_element("group-field"))
            row_elem.append(self._create_element("group-level", "1"))
            row_elem.append(self._create_element("next-row-new-page", "false"))
            row_elem.append(self._create_element("fixed-rows", "false"))
            row_elem.append(self._create_element("rows-per-page", "0"))
            row_elem.append(self._create_element("fill-rows", "false"))
            row_elem.append(self._create_element("group-separate", "false"))
            row_elem.append(self._create_element("take-first-row-only", "false"))
            
            row_data_bind2 = self._create_element("data-bind")
            row_data_bind2.set("_ref", data_bind_id)
            row_elem.append(row_data_bind2)
            
            cells_elem = self._create_element("cells")
            
            for cell_idx, cell in enumerate(row.cells):
                cell_elem = self._create_element("cell")
                cell_id = f"cell{self.id_counter + row_idx * len(row.cells) + cell_idx + 1}"
                cell_elem.append(self._create_element("id", cell_id))
                cell_elem.append(self._create_element("parent-id", row_id))
                cell_elem.append(self._create_element("z-index", "100"))
                cell_elem.append(self._create_element("visible", "true"))
                cell_elem.append(self._create_element("rotate", "0"))
                cell_elem.append(self._create_element("appear-in-every-page", "false"))
                cell_elem.append(self._create_element("cross-page", "false"))
                
                cell_bound = self._create_element("bound")
                cell_bound.append(self._create_element("x", "0"))
                cell_bound.append(self._create_element("y", "0"))
                cell_bound.append(self._create_element("width", "0"))
                cell_bound.append(self._create_element("height", "0"))
                cell_elem.append(cell_bound)
                
                cell_data_bind = self._create_element("data-bind")
                cell_data_bind.set("_ref", data_bind_id)
                cell_elem.append(cell_data_bind)
                
                cell_border = self._create_element("border")
                cell_border.set("_ref", border_id)
                cell_elem.append(cell_border)
                
                cell_bg = self._create_element("background")
                cell_bg.set("_ref", bg_id)
                cell_elem.append(cell_bg)
                
                cell_padding = self._create_element("padding")
                cell_padding.set("_ref", padding_id)
                cell_elem.append(cell_padding)
                
                cell_font = self._create_element("font")
                cell_font.set("_ref", font_id)
                cell_elem.append(cell_font)
                
                cell_para = self._create_element("paragraph")
                cell_para.set("_ref", para_id)
                cell_elem.append(cell_para)
                
                cell_format = self._create_element("format")
                cell_format.set("_ref", format_id)
                cell_elem.append(cell_format)
                
                cell_elem.append(self._create_element("shrink", "false"))
                
                children = self._create_element("children")
                if cell.content:
                    text_child = self._create_element("text")
                    text_child.append(self._create_element("id", f"text{cell_id}"))
                    text_child.append(self._create_element("parent-id", cell_id))
                    text_child.append(self._create_element("z-index", "100"))
                    text_child.append(self._create_element("visible", "true"))
                    text_child.append(self._create_element("rotate", "0"))
                    text_child.append(self._create_element("appear-in-every-page", "false"))

                    text_bound = self._create_element("bound")
                    text_bound.append(self._create_element("x", "0"))
                    text_bound.append(self._create_element("y", "0"))
                    text_bound.append(self._create_element("width", "0"))
                    text_bound.append(self._create_element("height", "0"))
                    text_child.append(text_bound)

                    cell_text_data_bind = self._create_element("data-bind")
                    cell_text_data_bind.set("_id", self._generate_id("DB"))
                    cell_text_bind = self._create_element("bind")
                    cell_text_bind.text = cell.content
                    cell_text_data_bind.append(cell_text_bind)
                    cell_text_data_bind.append(self._create_element("prepare-script"))
                    cell_text_data_bind.append(self._create_element("post-script"))
                    cell_text_data_bind.append(self._create_element("final-script"))
                    text_child.append(cell_text_data_bind)

                    text_border = self._create_element("border")
                    text_border.set("_ref", border_id)
                    text_child.append(text_border)

                    text_bg = self._create_element("background")
                    text_bg.set("_ref", bg_id)
                    text_child.append(text_bg)

                    text_padding = self._create_element("padding")
                    text_padding.set("_ref", padding_id)
                    text_child.append(text_padding)

                    text_font = self._create_element("font")
                    text_font.set("_ref", font_id)
                    text_child.append(text_font)

                    text_para = self._create_element("paragraph")
                    text_para.set("_ref", para_id)
                    text_child.append(text_para)

                    text_format = self._create_element("format")
                    text_format.set("_ref", format_id)
                    text_child.append(text_format)

                    text_child.append(self._create_element("text-type", "TEXT"))
                    text_child.append(self._create_element("use-rich-text-editor", "false"))

                    children.append(text_child)

                if cell.images:
                    image_child = self._create_element("image")
                    image_child.append(self._create_element("id", f"image{cell_id}"))
                    image_child.append(self._create_element("parent-id", cell_id))
                    image_child.append(self._create_element("z-index", "100"))
                    image_child.append(self._create_element("visible", "true"))
                    image_child.append(self._create_element("rotate", "0"))
                    image_child.append(self._create_element("appear-in-every-page", "false"))

                    image_bound = self._create_element("bound")
                    image_bound.append(self._create_element("x", "0"))
                    image_bound.append(self._create_element("y", "0"))
                    image_bound.append(self._create_element("width", "331"))
                    image_bound.append(self._create_element("height", "100"))
                    image_child.append(image_bound)

                    image_data_bind = self._create_element("data-bind")
                    image_data_bind.set("_id", self._generate_id("DB"))
                    image_bind = self._create_element("bind")
                    image_bind.text = cell.images[0]
                    image_data_bind.append(image_bind)
                    image_data_bind.append(self._create_element("prepare-script"))
                    image_data_bind.append(self._create_element("post-script"))
                    image_data_bind.append(self._create_element("final-script"))
                    image_child.append(image_data_bind)

                    image_border = self._create_element("border")
                    image_border.set("_ref", border_id)
                    image_child.append(image_border)

                    image_bg = self._create_element("background")
                    image_bg.set("_ref", bg_id)
                    image_child.append(image_bg)

                    image_padding = self._create_element("padding")
                    image_padding.set("_id", self._generate_id("PD"))
                    image_padding.append(self._create_element("left", "0"))
                    image_padding.append(self._create_element("right", "0"))
                    image_padding.append(self._create_element("top", "0"))
                    image_padding.append(self._create_element("bottom", "0"))
                    image_child.append(image_padding)

                    image_child.append(self._create_element("scale-type", "FIT"))
                    image_child.append(self._create_element("horizontal-position", "CENTER"))
                    image_child.append(self._create_element("vertical-position", "MIDDLE"))
                    image_child.append(self._create_element("cross-page-seal", "false"))
                    image_child.append(self._create_element("cross-page-seal-pages-type", "ALL"))
                    image_child.append(self._create_element("cache-image-in-execute-context", "true"))
                    image_child.append(self._create_element("reuse-image-in-pdf", "true"))
                    image_child.append(self._create_element("convert-images-to-jpeg-in-pdf", "false"))
                    image_child.append(self._create_element("extend", "false"))
                    image_child.append(self._create_element("extend-horizontal-gap"))
                    image_child.append(self._create_element("extend-vertical-gap"))

                    children.append(image_child)
                
                cell_elem.append(children)
                cell_elem.append(self._create_element("parent-cell"))
                # 根据内容类型设置单元格类型
                cell_type = "CONTAINER"
                if "<html>" in cell.content or "<body>" in cell.content:
                    cell_type = "HTML"
                elif cell.content:
                    cell_type = "TEXT"
                cell_elem.append(self._create_element("cell-type", cell_type))
                cell_elem.append(self._create_element("col-span", str(cell.col_span)))
                cell_elem.append(self._create_element("row-span", str(max(1, cell.row_span))))
                cell_elem.append(self._create_element("merged", "false"))
                cell_elem.append(self._create_element("use-rich-text-editor", "false"))
                cell_elem.append(self._create_element("overflow-hidden", "false"))
                cell_elem.append(self._create_element("merge-same-value-cells", "false"))
                cell_elem.append(self._create_element("merge-same-value-expression"))
                cell_elem.append(self._create_element("repeat-on-cross-page", "false"))
                
                cells_elem.append(cell_elem)
            
            row_elem.append(cells_elem)
            rows_elem.append(row_elem)
        
        grid.append(rows_elem)
        
        return grid

    def generate(self, page_content: PageContent, page_settings: dict) -> str:
        template = self._create_element("template")
        template.append(self._create_element("version", "2.0"))
        
        main_data_bind = self._create_data_bind()
        template.append(main_data_bind)
        main_data_bind_id = main_data_bind.get("_id")
        
        pages = self._create_element("pages")
        page = self._create_element("page")
        page.append(self._create_element("id", "page1"))
        
        labscare_ds = self._create_element("labscare-datasource")
        labscare_ds.append(self._create_element("template-type"))
        labscare_ds.append(self._create_element("case-group"))
        labscare_ds.append(self._create_element("template-id"))
        page.append(labscare_ds)
        
        paper = self._create_element("paper")
        paper.append(self._create_element("alias", "A4"))
        paper.append(self._create_element("width", str(page_settings['width'])))
        paper.append(self._create_element("height", str(page_settings['height'])))
        page.append(paper)
        
        design_region = self._create_element("design-region")
        design_region.append(self._create_element("width", str(page_settings['width'])))
        design_region.append(self._create_element("height", str(page_settings['height'])))
        design_region.append(self._create_element("synchronize-paper-width", "true"))
        design_region.append(self._create_element("synchronize-paper-height", "true"))
        page.append(design_region)
        
        page.append(self._create_element("paper-orientation", page_settings['orientation']))
        
        page_padding = self._create_element("page-padding")
        page_padding.append(self._create_element("left", str(page_settings['left_margin'])))
        page_padding.append(self._create_element("right", str(page_settings['right_margin'])))
        page_padding.append(self._create_element("top", str(page_settings['top_margin'])))
        page_padding.append(self._create_element("bottom", str(page_settings['bottom_margin'])))
        page.append(page_padding)
        
        economic_print = self._create_element("economic-print")
        economic_print.append(self._create_element("horizontal-economic", "false"))
        economic_print.append(self._create_element("vertical-economic", "false"))
        economic_print.append(self._create_element("horizontal-gap", "50"))
        economic_print.append(self._create_element("vertical-gap", "50"))
        page.append(economic_print)
        
        # 检查是否有背景图片
        has_background = False
        background_image = ""
        
        # 这里可以添加逻辑来检测文档中的背景图片
        # 暂时使用系统导出文件中的背景图片作为示例
        # 实际应用中，可以从文档中提取背景图片或允许用户指定
        
        if has_background:
            page_bg = self._create_element("background")
            page_bg.set("_id", self._generate_id("BG"))
            page_bg.append(self._create_element("type", "IMAGE"))
            page_bg.append(self._create_element("image", background_image))
            page_bg.append(self._create_element("repeat", "false"))
            page_bg.append(self._create_element("stretch", "true"))
            page_bg.append(self._create_element("print", "true"))
            page_bg.append(self._create_element("dpi", "255"))
            page_bg.append(self._create_element("x-offset", "0"))
            page_bg.append(self._create_element("y-offset", "0"))
        else:
            page_bg = self._create_background()
        page.append(page_bg)
        page_bg_id = page_bg.get("_id")
        
        page_data_bind = self._create_data_bind()
        page.append(page_data_bind)
        page_data_bind_id = page_data_bind.get("_id")
        
        border = self._create_border()
        border_id = border.get("_id")
        
        font = self._create_font()
        font_id = font.get("_id")
        
        para = self._create_paragraph()
        para_id = para.get("_id")
        
        format_elem = self._create_format()
        format_id = format_elem.get("_id")
        
        padding = self._create_padding()
        padding_id = padding.get("_id")
        
        header = self._create_element("header")
        header.append(self._create_element("id", "Header"))
        header.append(self._create_element("parent-id", "page1"))
        header.append(self._create_element("z-index", "100"))
        header.append(self._create_element("visible", "true"))
        header.append(self._create_element("rotate", "0"))
        
        header_data_bind = self._create_element("data-bind")
        header_data_bind.set("_ref", page_data_bind_id)
        header.append(header_data_bind)
        
        header_border = self._create_element("border")
        header_border.set("_ref", border_id)
        header.append(header_border)
        
        header_bg = self._create_element("background")
        header_bg.set("_ref", page_bg_id)
        header.append(header_bg)
        
        header_height = 298
        header.append(self._create_element("height", str(header_height)))
        header.append(self._create_element("shrink", "false"))
        
        header_children = self._create_element("children")
        
        if page_content.header:
            for item in page_content.header:
                if isinstance(item, TextContent):
                    text_elem = self._create_text_element(
                        item, page_data_bind_id, border_id, page_bg_id,
                        font_id, para_id, format_id, padding_id
                    )
                    text_elem.find("parent-id").text = "Header"
                    header_children.append(text_elem)
                elif isinstance(item, TableContent):
                    table_elem = self._create_table_element(
                        item, page_data_bind_id, border_id, page_bg_id,
                        font_id, para_id, format_id, padding_id
                    )
                    table_elem.find("parent-id").text = "Header"
                    header_children.append(table_elem)
        else:
            grid = self._create_element("grid")
            grid_id = f"grid{self.id_counter}"
            grid.append(self._create_element("id", grid_id))
            grid.append(self._create_element("parent-id", "Header"))
            grid.append(self._create_element("z-index", "100"))
            grid.append(self._create_element("visible", "true"))
            grid.append(self._create_element("appear-in-every-page", "false"))
            
            bound = self._create_element("bound")
            bound.append(self._create_element("x", "0"))
            bound.append(self._create_element("y", "0"))
            bound.append(self._create_element("width", "1782"))
            bound.append(self._create_element("height", "298"))
            grid.append(bound)
            
            data_bind_ref = self._create_element("data-bind")
            data_bind_ref.set("_ref", page_data_bind_id)
            grid.append(data_bind_ref)
            
            bg_ref = self._create_element("background")
            bg_ref.set("_ref", page_bg_id)
            grid.append(bg_ref)
            
            col_widths = self._create_element("column-widths")
            col_widths.text = "715.5,715.5,351"
            grid.append(col_widths)
            
            grid.append(self._create_element("row-colors"))
            grid.append(self._create_element("col-colors"))
            
            rows_elem = self._create_element("rows")
            
            row1 = self._create_element("row")
            row1.append(self._create_element("id", f"row{self.id_counter + 1}"))
            row1.append(self._create_element("parent-id", grid_id))
            row1.append(self._create_element("z-index", "100"))
            row1.append(self._create_element("visible", "true"))
            row1.append(self._create_element("rotate", "0"))
            row1.append(self._create_element("adjust-height", "true"))
            row1.append(self._create_element("appear-in-every-page", "false"))
            
            row_bound = self._create_element("bound")
            row_bound.append(self._create_element("x", "0"))
            row_bound.append(self._create_element("y", "0"))
            row_bound.append(self._create_element("width", "0"))
            row_bound.append(self._create_element("height", "0"))
            row1.append(row_bound)
            
            row_data_bind = self._create_element("data-bind")
            row_data_bind.set("_ref", page_data_bind_id)
            row1.append(row_data_bind)
            
            row_border = self._create_element("border")
            row_border.set("_ref", border_id)
            row1.append(row_border)
            
            row_bg = self._create_element("background")
            row_bg.set("_ref", page_bg_id)
            row1.append(row_bg)
            
            row1.append(self._create_element("enable-min-height", "false"))
            row1.append(self._create_element("min-height"))
            row1.append(self._create_element("height", "33.33333333333334"))
            row1.append(self._create_element("adjust-height", "true"))
            row1.append(self._create_element("appear-in-every-page", "false"))
            row1.append(self._create_element("row-type", "SIMPLE"))
            row1.append(self._create_element("group-field"))
            row1.append(self._create_element("group-level", "1"))
            row1.append(self._create_element("next-row-new-page", "false"))
            row1.append(self._create_element("fixed-rows", "false"))
            row1.append(self._create_element("rows-per-page", "0"))
            row1.append(self._create_element("fill-rows", "false"))
            row1.append(self._create_element("group-separate", "false"))
            row1.append(self._create_element("take-first-row-only", "false"))
            
            row_data_bind2 = self._create_element("data-bind")
            row_data_bind2.set("_ref", page_data_bind_id)
            row1.append(row_data_bind2)
            
            cells_elem = self._create_element("cells")
            
            cell1 = self._create_element("cell")
            cell1.append(self._create_element("id", f"cell{self.id_counter + 2}"))
            cell1.append(self._create_element("parent-id", f"row{self.id_counter + 1}"))
            cell1.append(self._create_element("z-index", "100"))
            cell1.append(self._create_element("visible", "true"))
            cell1.append(self._create_element("rotate", "0"))
            cell1.append(self._create_element("appear-in-every-page", "false"))
            cell1.append(self._create_element("cross-page", "false"))
            
            cell_bound = self._create_element("bound")
            cell_bound.append(self._create_element("x", "0"))
            cell_bound.append(self._create_element("y", "0"))
            cell_bound.append(self._create_element("width", "0"))
            cell_bound.append(self._create_element("height", "0"))
            cell1.append(cell_bound)
            
            cell_data_bind = self._create_element("data-bind")
            cell_data_bind.set("_ref", page_data_bind_id)
            cell1.append(cell_data_bind)
            
            cell_border = self._create_element("border")
            cell_border.set("_ref", border_id)
            cell1.append(cell_border)
            
            cell_bg = self._create_element("background")
            cell_bg.set("_ref", page_bg_id)
            cell1.append(cell_bg)
            
            cell_padding = self._create_element("padding")
            cell_padding.set("_ref", padding_id)
            cell1.append(cell_padding)
            
            cell_font = self._create_element("font")
            cell_font.set("_ref", font_id)
            cell1.append(cell_font)
            
            cell_para = self._create_element("paragraph")
            cell_para.set("_ref", para_id)
            cell1.append(cell_para)
            
            cell_format = self._create_element("format")
            cell_format.set("_ref", format_id)
            cell1.append(cell_format)
            
            cell1.append(self._create_element("shrink", "false"))
            cell1.append(self._create_element("children"))
            cell1.append(self._create_element("parent-cell"))
            cell1.append(self._create_element("cell-type", "CONTAINER"))
            cell1.append(self._create_element("col-span", "2"))
            cell1.append(self._create_element("row-span", "1"))
            cell1.append(self._create_element("merged", "false"))
            cell1.append(self._create_element("use-rich-text-editor", "false"))
            cell1.append(self._create_element("overflow-hidden", "false"))
            cell1.append(self._create_element("merge-same-value-cells", "false"))
            cell1.append(self._create_element("merge-same-value-expression"))
            cell1.append(self._create_element("repeat-on-cross-page", "false"))
            
            cells_elem.append(cell1)
            
            cell2 = self._create_element("cell")
            cell2.append(self._create_element("id", f"cell{self.id_counter + 3}"))
            cell2.append(self._create_element("parent-id", f"row{self.id_counter + 1}"))
            cell2.append(self._create_element("z-index", "100"))
            cell2.append(self._create_element("visible", "true"))
            cell2.append(self._create_element("rotate", "0"))
            cell2.append(self._create_element("appear-in-every-page", "false"))
            cell2.append(self._create_element("cross-page", "false"))
            
            cell_bound2 = self._create_element("bound")
            cell_bound2.append(self._create_element("x", "0"))
            cell_bound2.append(self._create_element("y", "0"))
            cell_bound2.append(self._create_element("width", "0"))
            cell_bound2.append(self._create_element("height", "0"))
            cell2.append(cell_bound2)
            
            cell_data_bind2 = self._create_element("data-bind")
            cell_data_bind2.set("_ref", page_data_bind_id)
            cell2.append(cell_data_bind2)
            
            cell_border2 = self._create_element("border")
            cell_border2.set("_ref", border_id)
            cell2.append(cell_border2)
            
            cell_bg2 = self._create_element("background")
            cell_bg2.set("_ref", page_bg_id)
            cell2.append(cell_bg2)
            
            cell_padding2 = self._create_element("padding")
            cell_padding2.set("_ref", padding_id)
            cell2.append(cell_padding2)
            
            cell_font2 = self._create_element("font")
            cell_font2.set("_ref", font_id)
            cell2.append(cell_font2)
            
            cell_para2 = self._create_element("paragraph")
            cell_para2.set("_ref", para_id)
            cell2.append(cell_para2)
            
            cell_format2 = self._create_element("format")
            cell_format2.set("_ref", format_id)
            cell2.append(cell_format2)
            
            cell2.append(self._create_element("shrink", "false"))
            cell2.append(self._create_element("children"))
            cell2.append(self._create_element("parent-cell"))
            cell2.append(self._create_element("cell-type", "HTML"))
            cell2.append(self._create_element("col-span", "1"))
            cell2.append(self._create_element("row-span", "1"))
            cell2.append(self._create_element("merged", "false"))
            cell2.append(self._create_element("use-rich-text-editor", "false"))
            cell2.append(self._create_element("overflow-hidden", "false"))
            cell2.append(self._create_element("merge-same-value-cells", "false"))
            cell2.append(self._create_element("merge-same-value-expression"))
            cell2.append(self._create_element("repeat-on-cross-page", "false"))
            
            cells_elem.append(cell2)
            
            row1.append(cells_elem)
            rows_elem.append(row1)
            
            row2 = self._create_element("row")
            row2.append(self._create_element("id", f"row{self.id_counter + 4}"))
            row2.append(self._create_element("parent-id", grid_id))
            row2.append(self._create_element("z-index", "100"))
            row2.append(self._create_element("visible", "true"))
            row2.append(self._create_element("rotate", "0"))
            row2.append(self._create_element("adjust-height", "true"))
            row2.append(self._create_element("appear-in-every-page", "false"))
            
            row_bound2 = self._create_element("bound")
            row_bound2.append(self._create_element("x", "0"))
            row_bound2.append(self._create_element("y", "0"))
            row_bound2.append(self._create_element("width", "0"))
            row_bound2.append(self._create_element("height", "0"))
            row2.append(row_bound2)
            
            row_data_bind3 = self._create_element("data-bind")
            row_data_bind3.set("_ref", page_data_bind_id)
            row2.append(row_data_bind3)
            
            row_border2 = self._create_element("border")
            row_border2.set("_ref", border_id)
            row2.append(row_border2)
            
            row_bg2 = self._create_element("background")
            row_bg2.set("_ref", page_bg_id)
            row2.append(row_bg2)
            
            row2.append(self._create_element("enable-min-height", "false"))
            row2.append(self._create_element("min-height"))
            row2.append(self._create_element("height", "66.66666666666667"))
            row2.append(self._create_element("adjust-height", "true"))
            row2.append(self._create_element("appear-in-every-page", "false"))
            row2.append(self._create_element("row-type", "SIMPLE"))
            row2.append(self._create_element("group-field"))
            row2.append(self._create_element("group-level", "1"))
            row2.append(self._create_element("next-row-new-page", "false"))
            row2.append(self._create_element("fixed-rows", "false"))
            row2.append(self._create_element("rows-per-page", "0"))
            row2.append(self._create_element("fill-rows", "false"))
            row2.append(self._create_element("group-separate", "false"))
            row2.append(self._create_element("take-first-row-only", "false"))
            
            row_data_bind4 = self._create_element("data-bind")
            row_data_bind4.set("_ref", page_data_bind_id)
            row2.append(row_data_bind4)
            
            cells_elem2 = self._create_element("cells")
            
            cell3 = self._create_element("cell")
            cell3.append(self._create_element("id", f"cell{self.id_counter + 5}"))
            cell3.append(self._create_element("parent-id", f"row{self.id_counter + 4}"))
            cell3.append(self._create_element("z-index", "100"))
            cell3.append(self._create_element("visible", "true"))
            cell3.append(self._create_element("rotate", "0"))
            cell3.append(self._create_element("appear-in-every-page", "false"))
            cell3.append(self._create_element("cross-page", "false"))
            
            cell_bound3 = self._create_element("bound")
            cell_bound3.append(self._create_element("x", "0"))
            cell_bound3.append(self._create_element("y", "0"))
            cell_bound3.append(self._create_element("width", "0"))
            cell_bound3.append(self._create_element("height", "0"))
            cell3.append(cell_bound3)
            
            cell_data_bind3 = self._create_element("data-bind")
            cell_data_bind3.set("_ref", page_data_bind_id)
            cell3.append(cell_data_bind3)
            
            cell_border3 = self._create_element("border")
            cell_border3.set("_ref", border_id)
            cell3.append(cell_border3)
            
            cell_bg3 = self._create_element("background")
            cell_bg3.set("_ref", page_bg_id)
            cell3.append(cell_bg3)
            
            cell_padding3 = self._create_element("padding")
            cell_padding3.set("_ref", padding_id)
            cell3.append(cell_padding3)
            
            cell_font3 = self._create_element("font")
            cell_font3.set("_ref", font_id)
            cell3.append(cell_font3)
            
            cell_para3 = self._create_element("paragraph")
            cell_para3.set("_ref", para_id)
            cell3.append(cell_para3)
            
            cell_format3 = self._create_element("format")
            cell_format3.set("_ref", format_id)
            cell3.append(cell_format3)
            
            cell3.append(self._create_element("shrink", "false"))
            cell3_children = self._create_element("children")
            if page_content.header_images:
                header_img = self._create_element("image")
                header_img.append(self._create_element("id", f"image{self.id_counter + 5}"))
                header_img.append(self._create_element("parent-id", f"cell{self.id_counter + 5}"))
                header_img.append(self._create_element("z-index", "100"))
                header_img.append(self._create_element("visible", "true"))
                header_img.append(self._create_element("rotate", "0"))
                header_img.append(self._create_element("appear-in-every-page", "false"))

                hb = self._create_element("bound")
                hb.append(self._create_element("x", "0"))
                hb.append(self._create_element("y", "-33.33333333333333"))
                hb.append(self._create_element("width", "331"))
                hb.append(self._create_element("height", "100"))
                header_img.append(hb)

                hdb = self._create_element("data-bind")
                hdb.set("_id", self._generate_id("DB"))
                hbv = self._create_element("bind")
                hbv.text = page_content.header_images[0]
                hdb.append(hbv)
                hdb.append(self._create_element("prepare-script"))
                hdb.append(self._create_element("post-script"))
                hdb.append(self._create_element("final-script"))
                header_img.append(hdb)

                hbd = self._create_element("border")
                hbd.set("_ref", border_id)
                header_img.append(hbd)

                hbg = self._create_element("background")
                hbg.set("_ref", page_bg_id)
                header_img.append(hbg)

                hpd = self._create_element("padding")
                hpd.set("_id", self._generate_id("PD"))
                hpd.append(self._create_element("left", "0"))
                hpd.append(self._create_element("right", "0"))
                hpd.append(self._create_element("top", "0"))
                hpd.append(self._create_element("bottom", "0"))
                header_img.append(hpd)

                header_img.append(self._create_element("scale-type", "FIT"))
                header_img.append(self._create_element("horizontal-position", "CENTER"))
                header_img.append(self._create_element("vertical-position", "MIDDLE"))
                header_img.append(self._create_element("cross-page-seal", "false"))
                header_img.append(self._create_element("cross-page-seal-pages-type", "ALL"))
                header_img.append(self._create_element("cache-image-in-execute-context", "true"))
                header_img.append(self._create_element("reuse-image-in-pdf", "true"))
                header_img.append(self._create_element("convert-images-to-jpeg-in-pdf", "false"))
                header_img.append(self._create_element("extend", "false"))
                header_img.append(self._create_element("extend-horizontal-gap"))
                header_img.append(self._create_element("extend-vertical-gap"))
                cell3_children.append(header_img)

            cell3.append(cell3_children)
            cell3.append(self._create_element("parent-cell"))
            cell3.append(self._create_element("cell-type", "CONTAINER"))
            cell3.append(self._create_element("col-span", "2"))
            cell3.append(self._create_element("row-span", "1"))
            cell3.append(self._create_element("merged", "false"))
            cell3.append(self._create_element("use-rich-text-editor", "false"))
            cell3.append(self._create_element("overflow-hidden", "false"))
            cell3.append(self._create_element("merge-same-value-cells", "false"))
            cell3.append(self._create_element("merge-same-value-expression"))
            cell3.append(self._create_element("repeat-on-cross-page", "false"))
            
            cells_elem2.append(cell3)
            
            cell4 = self._create_element("cell")
            cell4.append(self._create_element("id", f"cell{self.id_counter + 6}"))
            cell4.append(self._create_element("parent-id", f"row{self.id_counter + 4}"))
            cell4.append(self._create_element("z-index", "100"))
            cell4.append(self._create_element("visible", "true"))
            cell4.append(self._create_element("rotate", "0"))
            cell4.append(self._create_element("appear-in-every-page", "false"))
            cell4.append(self._create_element("cross-page", "false"))
            
            cell_bound4 = self._create_element("bound")
            cell_bound4.append(self._create_element("x", "0"))
            cell_bound4.append(self._create_element("y", "0"))
            cell_bound4.append(self._create_element("width", "0"))
            cell_bound4.append(self._create_element("height", "0"))
            cell4.append(cell_bound4)
            
            cell_data_bind4 = self._create_element("data-bind")
            cell_data_bind4.set("_ref", page_data_bind_id)
            cell4.append(cell_data_bind4)
            
            cell_border4 = self._create_element("border")
            cell_border4.set("_ref", border_id)
            cell4.append(cell_border4)
            
            cell_bg4 = self._create_element("background")
            cell_bg4.set("_ref", page_bg_id)
            cell4.append(cell_bg4)
            
            cell_padding4 = self._create_element("padding")
            cell_padding4.set("_ref", padding_id)
            cell4.append(cell_padding4)
            
            cell_font4 = self._create_element("font")
            cell_font4.set("_ref", font_id)
            cell4.append(cell_font4)
            
            cell_para4 = self._create_element("paragraph")
            cell_para4.set("_ref", para_id)
            cell4.append(cell_para4)
            
            cell_format4 = self._create_element("format")
            cell_format4.set("_ref", format_id)
            cell4.append(cell_format4)
            
            cell4.append(self._create_element("shrink", "false"))
            cell4_children = self._create_element("children")
            if len(page_content.header_images) > 1:
                header_img2 = self._create_element("image")
                header_img2.append(self._create_element("id", f"image{self.id_counter + 6}"))
                header_img2.append(self._create_element("parent-id", f"cell{self.id_counter + 6}"))
                header_img2.append(self._create_element("z-index", "100"))
                header_img2.append(self._create_element("visible", "true"))
                header_img2.append(self._create_element("rotate", "0"))
                header_img2.append(self._create_element("appear-in-every-page", "false"))

                hb2 = self._create_element("bound")
                hb2.append(self._create_element("x", "0"))
                hb2.append(self._create_element("y", "-33.33333333333333"))
                hb2.append(self._create_element("width", "331"))
                hb2.append(self._create_element("height", "100"))
                header_img2.append(hb2)

                hdb2 = self._create_element("data-bind")
                hdb2.set("_id", self._generate_id("DB"))
                hbv2 = self._create_element("bind")
                hbv2.text = page_content.header_images[1]
                hdb2.append(hbv2)
                hdb2.append(self._create_element("prepare-script"))
                hdb2.append(self._create_element("post-script"))
                hdb2.append(self._create_element("final-script"))
                header_img2.append(hdb2)

                hbd2 = self._create_element("border")
                hbd2.set("_ref", border_id)
                header_img2.append(hbd2)

                hbg2 = self._create_element("background")
                hbg2.set("_ref", page_bg_id)
                header_img2.append(hbg2)

                hpd2 = self._create_element("padding")
                hpd2.set("_id", self._generate_id("PD"))
                hpd2.append(self._create_element("left", "0"))
                hpd2.append(self._create_element("right", "0"))
                hpd2.append(self._create_element("top", "0"))
                hpd2.append(self._create_element("bottom", "0"))
                header_img2.append(hpd2)

                header_img2.append(self._create_element("scale-type", "FIT"))
                header_img2.append(self._create_element("horizontal-position", "CENTER"))
                header_img2.append(self._create_element("vertical-position", "MIDDLE"))
                header_img2.append(self._create_element("cross-page-seal", "false"))
                header_img2.append(self._create_element("cross-page-seal-pages-type", "ALL"))
                header_img2.append(self._create_element("cache-image-in-execute-context", "true"))
                header_img2.append(self._create_element("reuse-image-in-pdf", "true"))
                header_img2.append(self._create_element("convert-images-to-jpeg-in-pdf", "false"))
                header_img2.append(self._create_element("extend", "false"))
                header_img2.append(self._create_element("extend-horizontal-gap"))
                header_img2.append(self._create_element("extend-vertical-gap"))
                cell4_children.append(header_img2)

            cell4.append(cell4_children)
            cell4.append(self._create_element("parent-cell"))
            cell4.append(self._create_element("cell-type", "HTML"))
            cell4.append(self._create_element("col-span", "1"))
            cell4.append(self._create_element("row-span", "1"))
            cell4.append(self._create_element("merged", "false"))
            cell4.append(self._create_element("use-rich-text-editor", "false"))
            cell4.append(self._create_element("overflow-hidden", "false"))
            cell4.append(self._create_element("merge-same-value-cells", "false"))
            cell4.append(self._create_element("merge-same-value-expression"))
            cell4.append(self._create_element("repeat-on-cross-page", "false"))
            
            cells_elem2.append(cell4)
            
            row2.append(cells_elem2)
            rows_elem.append(row2)
            
            row3 = self._create_element("row")
            row3.append(self._create_element("id", f"row{self.id_counter + 7}"))
            row3.append(self._create_element("parent-id", grid_id))
            row3.append(self._create_element("z-index", "100"))
            row3.append(self._create_element("visible", "true"))
            row3.append(self._create_element("rotate", "0"))
            row3.append(self._create_element("adjust-height", "true"))
            row3.append(self._create_element("appear-in-every-page", "false"))
            
            row_bound3 = self._create_element("bound")
            row_bound3.append(self._create_element("x", "0"))
            row_bound3.append(self._create_element("y", "0"))
            row_bound3.append(self._create_element("width", "0"))
            row_bound3.append(self._create_element("height", "0"))
            row3.append(row_bound3)
            
            row_data_bind5 = self._create_element("data-bind")
            row_data_bind5.set("_ref", page_data_bind_id)
            row3.append(row_data_bind5)
            
            row_border3 = self._create_element("border")
            row_border3.set("_ref", border_id)
            row3.append(row_border3)
            
            row_bg3 = self._create_element("background")
            row_bg3.set("_ref", page_bg_id)
            row3.append(row_bg3)
            
            row3.append(self._create_element("enable-min-height", "false"))
            row3.append(self._create_element("min-height"))
            row3.append(self._create_element("height", "66.66666666666667"))
            row3.append(self._create_element("adjust-height", "true"))
            row3.append(self._create_element("appear-in-every-page", "false"))
            row3.append(self._create_element("row-type", "SIMPLE"))
            row3.append(self._create_element("group-field"))
            row3.append(self._create_element("group-level", "1"))
            row3.append(self._create_element("next-row-new-page", "false"))
            row3.append(self._create_element("fixed-rows", "false"))
            row3.append(self._create_element("rows-per-page", "0"))
            row3.append(self._create_element("fill-rows", "false"))
            row3.append(self._create_element("group-separate", "false"))
            row3.append(self._create_element("take-first-row-only", "false"))
            
            row_data_bind6 = self._create_element("data-bind")
            row_data_bind6.set("_ref", page_data_bind_id)
            row3.append(row_data_bind6)
            
            cells_elem3 = self._create_element("cells")
            
            cell5 = self._create_element("cell")
            cell5.append(self._create_element("id", f"cell{self.id_counter + 8}"))
            cell5.append(self._create_element("parent-id", f"row{self.id_counter + 7}"))
            cell5.append(self._create_element("z-index", "100"))
            cell5.append(self._create_element("visible", "true"))
            cell5.append(self._create_element("rotate", "0"))
            cell5.append(self._create_element("appear-in-every-page", "false"))
            cell5.append(self._create_element("cross-page", "false"))
            
            cell_bound5 = self._create_element("bound")
            cell_bound5.append(self._create_element("x", "0"))
            cell_bound5.append(self._create_element("y", "0"))
            cell_bound5.append(self._create_element("width", "0"))
            cell_bound5.append(self._create_element("height", "0"))
            cell5.append(cell_bound5)
            
            cell_data_bind5 = self._create_element("data-bind")
            cell_data_bind5.set("_id", self._generate_id("DB"))
            cell_bind5 = self._create_element("bind")
            cell_bind5.text = "恒洁卫浴集团有限公司质量中心"
            cell_data_bind5.append(cell_bind5)
            cell_data_bind5.append(self._create_element("prepare-script"))
            cell_data_bind5.append(self._create_element("post-script"))
            cell_data_bind5.append(self._create_element("final-script"))
            cell5.append(cell_data_bind5)
            
            cell_border5 = self._create_element("border")
            cell_border5.set("_id", self._generate_id("BD"))
            cell_border5.append(self._create_element("left"))
            cell_border5.append(self._create_element("right"))
            top_border = self._create_element("top")
            line_type = self._create_element("line-type")
            aside = self._create_element("aside")
            middle = self._create_element("middle")
            middle.append(self._create_element("dash"))
            middle.append(self._create_element("width", "1.7638888888888888"))
            middle.append(self._create_element("dummy", "false"))
            bside = self._create_element("bside")
            line_type.append(aside)
            line_type.append(middle)
            line_type.append(bside)
            top_border.append(line_type)
            top_border.append(self._create_element("color", "rgba(0,0,0,1)"))
            cell_border5.append(top_border)
            cell_border5.append(self._create_element("bottom"))
            cell_border5.append(self._create_element("tr2bl"))
            cell_border5.append(self._create_element("tl2br"))
            cell5.append(cell_border5)
            
            cell_bg5 = self._create_element("background")
            cell_bg5.set("_ref", page_bg_id)
            cell5.append(cell_bg5)
            
            cell_padding5 = self._create_element("padding")
            cell_padding5.set("_ref", padding_id)
            cell5.append(cell_padding5)
            
            cell_font5 = self._create_element("font")
            cell_font5.set("_id", self._generate_id("FT"))
            cell_font5.append(self._create_element("chinese-font", "宋体"))
            cell_font5.append(self._create_element("western-font", "Times New Roman"))
            cell_font5.append(self._create_element("bold", "false"))
            cell_font5.append(self._create_element("italic", "false"))
            cell_font5.append(self._create_element("under-line", "false"))
            cell_font5.append(self._create_element("delete-line", "false"))
            cell_font5.append(self._create_element("over-line", "false"))
            cell_font5.append(self._create_element("under-line-style"))
            cell_font5.append(self._create_element("delete-line-style"))
            cell_font5.append(self._create_element("over-line-style"))
            font_size = self._create_element("font-size")
            font_size.append(self._create_element("alias", "五号"))
            font_size.append(self._create_element("pounds", "10.5"))
            cell_font5.append(font_size)
            cell_font5.append(self._create_element("color", "rgba(0,0,0,1)"))
            cell_font5.append(self._create_element("highlight", "false"))
            cell_font5.append(self._create_element("highlight-color", "rgba(255,255,255,1)"))
            cell5.append(cell_font5)
            
            cell_para5 = self._create_element("paragraph")
            cell_para5.set("_id", self._generate_id("PG"))
            cell_para5.append(self._create_element("writing-mode", "HORIZONTAL_TB"))
            cell_para5.append(self._create_element("text-align", "LEFT"))
            cell_para5.append(self._create_element("vertical-align", "MIDDLE"))
            cell_para5.append(self._create_element("text-indent", "0"))
            cell_para5.append(self._create_element("hanging-indent", "0"))
            cell_para5.append(self._create_element("line-height", "1"))
            cell_para5.append(self._create_element("line-height-unit", "LINE"))
            cell_para5.append(self._create_element("space-before", "0"))
            cell_para5.append(self._create_element("space-before-unit", "LINE"))
            cell_para5.append(self._create_element("space-after", "0"))
            cell_para5.append(self._create_element("space-after-unit", "LINE"))
            cell_para5.append(self._create_element("allow-breaking", "false"))
            cell_para5.append(self._create_element("shrink", "false"))
            cell_para5.append(self._create_element("letter-spacing", "0"))
            cell_para5.append(self._create_element("word-spacing", "0"))
            cell_para5.append(self._create_element("auto-space-dn", "false"))
            cell_para5.append(self._create_element("auto-space-de", "false"))
            cell_para5.append(self._create_element("line-wrap", "true"))
            cell_para5.append(self._create_element("cells-filling", "false"))
            cell_para5.append(self._create_element("cells", "0"))
            cell_line_style = self._create_element("cell-line-style")
            line_type2 = self._create_element("line-type")
            aside2 = self._create_element("aside")
            middle2 = self._create_element("middle")
            middle2.append(self._create_element("dash"))
            middle2.append(self._create_element("width", "1.7638888888888888"))
            middle2.append(self._create_element("dummy", "false"))
            bside2 = self._create_element("bside")
            line_type2.append(aside2)
            line_type2.append(middle2)
            line_type2.append(bside2)
            cell_line_style.append(line_type2)
            cell_line_style.append(self._create_element("color", "rgba(0,0,0,1)"))
            cell_para5.append(cell_line_style)
            cell5.append(cell_para5)
            
            cell_format5 = self._create_element("format")
            cell_format5.set("_ref", format_id)
            cell5.append(cell_format5)
            
            cell5.append(self._create_element("shrink", "false"))
            cell5.append(self._create_element("children"))
            cell5.append(self._create_element("parent-cell"))
            cell5.append(self._create_element("cell-type", "TEXT"))
            cell5.append(self._create_element("col-span", "1"))
            cell5.append(self._create_element("row-span", "1"))
            cell5.append(self._create_element("merged", "false"))
            cell5.append(self._create_element("use-rich-text-editor", "false"))
            cell5.append(self._create_element("overflow-hidden", "false"))
            cell5.append(self._create_element("merge-same-value-cells", "false"))
            cell5.append(self._create_element("merge-same-value-expression"))
            cell5.append(self._create_element("repeat-on-cross-page", "false"))
            
            cells_elem3.append(cell5)
            
            cell6 = self._create_element("cell")
            cell6.append(self._create_element("id", f"cell{self.id_counter + 9}"))
            cell6.append(self._create_element("parent-id", f"row{self.id_counter + 7}"))
            cell6.append(self._create_element("z-index", "100"))
            cell6.append(self._create_element("visible", "true"))
            cell6.append(self._create_element("rotate", "0"))
            cell6.append(self._create_element("appear-in-every-page", "false"))
            cell6.append(self._create_element("cross-page", "false"))
            
            cell_bound6 = self._create_element("bound")
            cell_bound6.append(self._create_element("x", "0"))
            cell_bound6.append(self._create_element("y", "0"))
            cell_bound6.append(self._create_element("width", "0"))
            cell_bound6.append(self._create_element("height", "0"))
            cell6.append(cell_bound6)
            
            cell_data_bind6 = self._create_element("data-bind")
            cell_data_bind6.set("_id", self._generate_id("DB"))
            cell_bind6 = self._create_element("bind")
            cell_bind6.text = "报告编号：    "
            cell_data_bind6.append(cell_bind6)
            cell_data_bind6.append(self._create_element("prepare-script"))
            cell_data_bind6.append(self._create_element("post-script"))
            cell_data_bind6.append(self._create_element("final-script"))
            cell6.append(cell_data_bind6)
            
            cell_border6 = self._create_element("border")
            cell_border6.set("_ref", cell_border5.get("_id"))
            cell6.append(cell_border6)
            
            cell_bg6 = self._create_element("background")
            cell_bg6.set("_ref", page_bg_id)
            cell6.append(cell_bg6)
            
            cell_padding6 = self._create_element("padding")
            cell_padding6.set("_ref", padding_id)
            cell6.append(cell_padding6)
            
            cell_font6 = self._create_element("font")
            cell_font6.set("_id", self._generate_id("FT"))
            cell_font6.append(self._create_element("chinese-font", "宋体"))
            cell_font6.append(self._create_element("western-font", "Times New Roman"))
            cell_font6.append(self._create_element("bold", "false"))
            cell_font6.append(self._create_element("italic", "false"))
            cell_font6.append(self._create_element("under-line", "false"))
            cell_font6.append(self._create_element("delete-line", "false"))
            cell_font6.append(self._create_element("over-line", "false"))
            cell_font6.append(self._create_element("under-line-style"))
            cell_font6.append(self._create_element("delete-line-style"))
            cell_font6.append(self._create_element("over-line-style"))
            font_size2 = self._create_element("font-size")
            font_size2.append(self._create_element("alias", "五号"))
            font_size2.append(self._create_element("pounds", "10.5"))
            cell_font6.append(font_size2)
            cell_font6.append(self._create_element("color", "rgba(0,0,0,1)"))
            cell_font6.append(self._create_element("highlight", "false"))
            cell_font6.append(self._create_element("highlight-color", "rgba(255,255,255,1)"))
            cell6.append(cell_font6)
            
            cell_para6 = self._create_element("paragraph")
            cell_para6.set("_id", self._generate_id("PG"))
            cell_para6.append(self._create_element("writing-mode", "HORIZONTAL_TB"))
            cell_para6.append(self._create_element("text-align", "RIGHT"))
            cell_para6.append(self._create_element("vertical-align", "MIDDLE"))
            cell_para6.append(self._create_element("text-indent", "0"))
            cell_para6.append(self._create_element("hanging-indent", "0"))
            cell_para6.append(self._create_element("line-height", "1"))
            cell_para6.append(self._create_element("line-height-unit", "LINE"))
            cell_para6.append(self._create_element("space-before", "0"))
            cell_para6.append(self._create_element("space-before-unit", "LINE"))
            cell_para6.append(self._create_element("space-after", "0"))
            cell_para6.append(self._create_element("space-after-unit", "LINE"))
            cell_para6.append(self._create_element("allow-breaking", "false"))
            cell_para6.append(self._create_element("shrink", "false"))
            cell_para6.append(self._create_element("letter-spacing", "0"))
            cell_para6.append(self._create_element("word-spacing", "0"))
            cell_para6.append(self._create_element("auto-space-dn", "false"))
            cell_para6.append(self._create_element("auto-space-de", "false"))
            cell_para6.append(self._create_element("line-wrap", "true"))
            cell_para6.append(self._create_element("cells-filling", "false"))
            cell_para6.append(self._create_element("cells", "0"))
            cell_line_style2 = self._create_element("cell-line-style")
            line_type3 = self._create_element("line-type")
            aside3 = self._create_element("aside")
            middle3 = self._create_element("middle")
            middle3.append(self._create_element("dash"))
            middle3.append(self._create_element("width", "1.7638888888888888"))
            middle3.append(self._create_element("dummy", "false"))
            bside3 = self._create_element("bside")
            line_type3.append(aside3)
            line_type3.append(middle3)
            line_type3.append(bside3)
            cell_line_style2.append(line_type3)
            cell_line_style2.append(self._create_element("color", "rgba(0,0,0,1)"))
            cell_para6.append(cell_line_style2)
            cell6.append(cell_para6)
            
            cell_format6 = self._create_element("format")
            cell_format6.set("_ref", format_id)
            cell6.append(cell_format6)
            
            cell6.append(self._create_element("shrink", "false"))
            cell6.append(self._create_element("children"))
            cell6.append(self._create_element("parent-cell"))
            cell6.append(self._create_element("cell-type", "TEXT"))
            cell6.append(self._create_element("col-span", "1"))
            cell6.append(self._create_element("row-span", "1"))
            cell6.append(self._create_element("merged", "false"))
            cell6.append(self._create_element("use-rich-text-editor", "false"))
            cell6.append(self._create_element("overflow-hidden", "false"))
            cell6.append(self._create_element("merge-same-value-cells", "false"))
            cell6.append(self._create_element("merge-same-value-expression"))
            cell6.append(self._create_element("repeat-on-cross-page", "false"))
            
            cells_elem3.append(cell6)
            
            cell7 = self._create_element("cell")
            cell7.append(self._create_element("id", f"cell{self.id_counter + 10}"))
            cell7.append(self._create_element("parent-id", f"row{self.id_counter + 7}"))
            cell7.append(self._create_element("z-index", "100"))
            cell7.append(self._create_element("visible", "true"))
            cell7.append(self._create_element("rotate", "0"))
            cell7.append(self._create_element("appear-in-every-page", "false"))
            cell7.append(self._create_element("cross-page", "false"))
            
            cell_bound7 = self._create_element("bound")
            cell_bound7.append(self._create_element("x", "0"))
            cell_bound7.append(self._create_element("y", "0"))
            cell_bound7.append(self._create_element("width", "0"))
            cell_bound7.append(self._create_element("height", "0"))
            cell7.append(cell_bound7)
            
            cell_data_bind7 = self._create_element("data-bind")
            cell_data_bind7.set("_id", self._generate_id("DB"))
            cell_bind7 = self._create_element("bind")
            cell_bind7.text = "  版本：A版"
            cell_data_bind7.append(cell_bind7)
            cell_data_bind7.append(self._create_element("prepare-script"))
            cell_data_bind7.append(self._create_element("post-script"))
            cell_data_bind7.append(self._create_element("final-script"))
            cell7.append(cell_data_bind7)
            
            cell_border7 = self._create_element("border")
            cell_border7.set("_ref", cell_border5.get("_id"))
            cell7.append(cell_border7)
            
            cell_bg7 = self._create_element("background")
            cell_bg7.set("_ref", page_bg_id)
            cell7.append(cell_bg7)
            
            cell_padding7 = self._create_element("padding")
            cell_padding7.set("_ref", padding_id)
            cell7.append(cell_padding7)
            
            cell_font7 = self._create_element("font")
            cell_font7.set("_id", self._generate_id("FT"))
            cell_font7.append(self._create_element("chinese-font", "宋体"))
            cell_font7.append(self._create_element("western-font", "Times New Roman"))
            cell_font7.append(self._create_element("bold", "false"))
            cell_font7.append(self._create_element("italic", "false"))
            cell_font7.append(self._create_element("under-line", "false"))
            cell_font7.append(self._create_element("delete-line", "false"))
            cell_font7.append(self._create_element("over-line", "false"))
            cell_font7.append(self._create_element("under-line-style"))
            cell_font7.append(self._create_element("delete-line-style"))
            cell_font7.append(self._create_element("over-line-style"))
            font_size3 = self._create_element("font-size")
            font_size3.append(self._create_element("alias", "小四"))
            font_size3.append(self._create_element("pounds", "12"))
            cell_font7.append(font_size3)
            cell_font7.append(self._create_element("color", "rgba(0,0,0,1)"))
            cell_font7.append(self._create_element("highlight", "false"))
            cell_font7.append(self._create_element("highlight-color", "rgba(255,255,255,1)"))
            cell7.append(cell_font7)
            
            cell_para7 = self._create_element("paragraph")
            cell_para7.set("_ref", para_id)
            cell7.append(cell_para7)
            
            cell_format7 = self._create_element("format")
            cell_format7.set("_ref", format_id)
            cell7.append(cell_format7)
            
            cell7.append(self._create_element("shrink", "false"))
            cell7.append(self._create_element("children"))
            cell7.append(self._create_element("parent-cell"))
            cell7.append(self._create_element("cell-type", "HTML"))
            cell7.append(self._create_element("col-span", "1"))
            cell7.append(self._create_element("row-span", "1"))
            cell7.append(self._create_element("merged", "false"))
            cell7.append(self._create_element("use-rich-text-editor", "false"))
            cell7.append(self._create_element("overflow-hidden", "false"))
            cell7.append(self._create_element("merge-same-value-cells", "false"))
            cell7.append(self._create_element("merge-same-value-expression"))
            cell7.append(self._create_element("repeat-on-cross-page", "false"))
            
            cells_elem3.append(cell7)
            
            row3.append(cells_elem3)
            rows_elem.append(row3)
            
            row4 = self._create_element("row")
            row4.append(self._create_element("id", f"row{self.id_counter + 11}"))
            row4.append(self._create_element("parent-id", grid_id))
            row4.append(self._create_element("z-index", "100"))
            row4.append(self._create_element("visible", "true"))
            row4.append(self._create_element("rotate", "0"))
            row4.append(self._create_element("adjust-height", "true"))
            row4.append(self._create_element("appear-in-every-page", "false"))
            
            row_bound4 = self._create_element("bound")
            row_bound4.append(self._create_element("x", "0"))
            row_bound4.append(self._create_element("y", "0"))
            row_bound4.append(self._create_element("width", "0"))
            row_bound4.append(self._create_element("height", "0"))
            row4.append(row_bound4)
            
            row_data_bind7 = self._create_element("data-bind")
            row_data_bind7.set("_ref", page_data_bind_id)
            row4.append(row_data_bind7)
            
            row_border4 = self._create_element("border")
            row_border4.set("_ref", border_id)
            row4.append(row_border4)
            
            row_bg4 = self._create_element("background")
            row_bg4.set("_ref", page_bg_id)
            row4.append(row_bg4)
            
            row4.append(self._create_element("enable-min-height", "false"))
            row4.append(self._create_element("min-height"))
            row4.append(self._create_element("height", "131.33333333333331"))
            row4.append(self._create_element("adjust-height", "true"))
            row4.append(self._create_element("appear-in-every-page", "false"))
            row4.append(self._create_element("row-type", "SIMPLE"))
            row4.append(self._create_element("group-field"))
            row4.append(self._create_element("group-level", "1"))
            row4.append(self._create_element("next-row-new-page", "false"))
            row4.append(self._create_element("fixed-rows", "false"))
            row4.append(self._create_element("rows-per-page", "0"))
            row4.append(self._create_element("fill-rows", "false"))
            row4.append(self._create_element("group-separate", "false"))
            row4.append(self._create_element("take-first-row-only", "false"))
            
            row_data_bind8 = self._create_element("data-bind")
            row_data_bind8.set("_ref", page_data_bind_id)
            row4.append(row_data_bind8)
            
            cells_elem4 = self._create_element("cells")
            
            cell8 = self._create_element("cell")
            cell8.append(self._create_element("id", f"cell{self.id_counter + 12}"))
            cell8.append(self._create_element("parent-id", f"row{self.id_counter + 11}"))
            cell8.append(self._create_element("z-index", "100"))
            cell8.append(self._create_element("visible", "true"))
            cell8.append(self._create_element("rotate", "0"))
            cell8.append(self._create_element("appear-in-every-page", "false"))
            cell8.append(self._create_element("cross-page", "false"))
            
            cell_bound8 = self._create_element("bound")
            cell_bound8.append(self._create_element("x", "0"))
            cell_bound8.append(self._create_element("y", "0"))
            cell_bound8.append(self._create_element("width", "0"))
            cell_bound8.append(self._create_element("height", "0"))
            cell8.append(cell_bound8)
            
            cell_data_bind8 = self._create_element("data-bind")
            cell_data_bind8.set("_id", self._generate_id("DB"))
            cell_bind8 = self._create_element("bind")
            cell_bind8.text = "检 验 报 告 "
            cell_data_bind8.append(cell_bind8)
            cell_data_bind8.append(self._create_element("prepare-script"))
            cell_data_bind8.append(self._create_element("post-script"))
            cell_data_bind8.append(self._create_element("final-script"))
            cell8.append(cell_data_bind8)
            
            cell_border8 = self._create_element("border")
            cell_border8.set("_ref", border_id)
            cell8.append(cell_border8)
            
            cell_bg8 = self._create_element("background")
            cell_bg8.set("_ref", page_bg_id)
            cell8.append(cell_bg8)
            
            cell_padding8 = self._create_element("padding")
            cell_padding8.set("_ref", padding_id)
            cell8.append(cell_padding8)
            
            cell_font8 = self._create_element("font")
            cell_font8.set("_id", self._generate_id("FT"))
            cell_font8.append(self._create_element("chinese-font", "新宋体"))
            cell_font8.append(self._create_element("western-font", "Times New Roman"))
            cell_font8.append(self._create_element("bold", "true"))
            cell_font8.append(self._create_element("italic", "false"))
            cell_font8.append(self._create_element("under-line", "false"))
            cell_font8.append(self._create_element("delete-line", "false"))
            cell_font8.append(self._create_element("over-line", "false"))
            cell_font8.append(self._create_element("under-line-style"))
            cell_font8.append(self._create_element("delete-line-style"))
            cell_font8.append(self._create_element("over-line-style"))
            font_size4 = self._create_element("font-size")
            font_size4.append(self._create_element("alias", "二号"))
            font_size4.append(self._create_element("pounds", "22"))
            cell_font8.append(font_size4)
            cell_font8.append(self._create_element("color", "rgba(0,0,0,1)"))
            cell_font8.append(self._create_element("highlight", "false"))
            cell_font8.append(self._create_element("highlight-color", "rgba(255,255,255,1)"))
            cell8.append(cell_font8)
            
            cell_para8 = self._create_element("paragraph")
            cell_para8.set("_ref", para_id)
            cell8.append(cell_para8)
            
            cell_format8 = self._create_element("format")
            cell_format8.set("_ref", format_id)
            cell8.append(cell_format8)
            
            cell8.append(self._create_element("shrink", "false"))
            cell8.append(self._create_element("children"))
            cell8.append(self._create_element("parent-cell"))
            cell8.append(self._create_element("cell-type", "TEXT"))
            cell8.append(self._create_element("col-span", "3"))
            cell8.append(self._create_element("row-span", "1"))
            cell8.append(self._create_element("merged", "false"))
            cell8.append(self._create_element("use-rich-text-editor", "false"))
            cell8.append(self._create_element("overflow-hidden", "false"))
            cell8.append(self._create_element("merge-same-value-cells", "false"))
            cell8.append(self._create_element("merge-same-value-expression"))
            cell8.append(self._create_element("repeat-on-cross-page", "false"))
            
            cells_elem4.append(cell8)
            
            row4.append(cells_elem4)
            rows_elem.append(row4)
            
            grid.append(rows_elem)
            
            header_children.append(grid)
        
        header.append(header_children)
        page.append(header)
        
        body = self._create_element("body")
        body.append(self._create_element("id", "Body"))
        body.append(self._create_element("parent-id", "page1"))
        body.append(self._create_element("z-index", "100"))
        body.append(self._create_element("visible", "true"))
        body.append(self._create_element("rotate", "0"))
        
        body_data_bind = self._create_element("data-bind")
        body_data_bind.set("_ref", page_data_bind_id)
        body.append(body_data_bind)
        
        body_border = self._create_element("border")
        body_border.set("_ref", border_id)
        body.append(body_border)
        
        body_bg = self._create_element("background")
        body_bg.set("_ref", page_bg_id)
        body.append(body_bg)
        
        body.append(self._create_element("shrink", "false"))
        
        body_children = self._create_element("children")
        y_offset = 0
        for item in page_content.body:
            if isinstance(item, TextContent):
                item.y = y_offset
                text_elem = self._create_text_element(
                    item, page_data_bind_id, border_id, page_bg_id,
                    font_id, para_id, format_id, padding_id
                )
                body_children.append(text_elem)
                y_offset += 35
            elif isinstance(item, TableContent):
                item.y = y_offset
                table_elem = self._create_table_element(
                    item, page_data_bind_id, border_id, page_bg_id,
                    font_id, para_id, format_id, padding_id
                )
                body_children.append(table_elem)
                row_count = len(item.rows)
                y_offset += row_count * 35
        body.append(body_children)
        page.append(body)
        
        footer = self._create_element("footer")
        footer.append(self._create_element("id", "Footer"))
        footer.append(self._create_element("parent-id", "page1"))
        footer.append(self._create_element("z-index", "100"))
        footer.append(self._create_element("visible", "true"))
        footer.append(self._create_element("rotate", "0"))
        
        footer_data_bind = self._create_element("data-bind")
        footer_data_bind.set("_ref", page_data_bind_id)
        footer.append(footer_data_bind)
        
        footer_border = self._create_element("border")
        footer_border.set("_ref", border_id)
        footer.append(footer_border)
        
        footer_bg = self._create_element("background")
        footer_bg.set("_ref", page_bg_id)
        footer.append(footer_bg)
        
        footer_height = 100
        if page_content.footer:
            footer_height = max(100, len(page_content.footer) * 35)
        footer.append(self._create_element("height", str(footer_height)))
        footer.append(self._create_element("shrink", "false"))
        
        footer_children = self._create_element("children")
        for item in page_content.footer:
            if isinstance(item, TextContent):
                text_elem = self._create_text_element(
                    item, page_data_bind_id, border_id, page_bg_id,
                    font_id, para_id, format_id, padding_id
                )
                text_elem.find("parent-id").text = "Footer"
                footer_children.append(text_elem)
            elif isinstance(item, TableContent):
                table_elem = self._create_table_element(
                    item, page_data_bind_id, border_id, page_bg_id,
                    font_id, para_id, format_id, padding_id
                )
                table_elem.find("parent-id").text = "Footer"
                footer_children.append(table_elem)
        footer.append(footer_children)
        page.append(footer)
        
        pages.append(page)

        # 追加第二页脚本页，匹配目标xmreport的两页结构
        page2 = self._create_element("page")
        page2.append(self._create_element("id", "page"))

        labscare_ds2 = self._create_element("labscare-datasource")
        labscare_ds2.append(self._create_element("template-type"))
        labscare_ds2.append(self._create_element("case-group"))
        labscare_ds2.append(self._create_element("template-id"))
        page2.append(labscare_ds2)

        paper2 = self._create_element("paper")
        paper2.append(self._create_element("alias", "A4"))
        paper2.append(self._create_element("width", str(page_settings['width'])))
        paper2.append(self._create_element("height", str(page_settings['height'])))
        page2.append(paper2)

        design_region2 = self._create_element("design-region")
        design_region2.append(self._create_element("width", str(page_settings['width'])))
        design_region2.append(self._create_element("height", str(page_settings['height'])))
        design_region2.append(self._create_element("synchronize-paper-width", "true"))
        design_region2.append(self._create_element("synchronize-paper-height", "true"))
        page2.append(design_region2)

        page2.append(self._create_element("paper-orientation", page_settings['orientation']))

        page_padding2 = self._create_element("page-padding")
        page_padding2.append(self._create_element("left", str(page_settings['left_margin'])))
        page_padding2.append(self._create_element("right", str(page_settings['right_margin'])))
        page_padding2.append(self._create_element("top", str(page_settings['top_margin'])))
        page_padding2.append(self._create_element("bottom", str(page_settings['bottom_margin'])))
        page2.append(page_padding2)

        economic_print2 = self._create_element("economic-print")
        economic_print2.append(self._create_element("horizontal-economic", "false"))
        economic_print2.append(self._create_element("vertical-economic", "false"))
        economic_print2.append(self._create_element("horizontal-gap", "50"))
        economic_print2.append(self._create_element("vertical-gap", "50"))
        page2.append(economic_print2)

        page2_bg = self._create_element("background")
        page2_bg.set("_ref", page_bg_id)
        page2.append(page2_bg)

        page2_db = self._create_element("data-bind")
        page2_db.set("_ref", page_data_bind_id)
        page2.append(page2_db)

        header2 = self._create_element("header")
        header2.append(self._create_element("id", "Header"))
        header2.append(self._create_element("parent-id", "page"))
        header2.append(self._create_element("z-index", "100"))
        header2.append(self._create_element("visible", "true"))
        header2.append(self._create_element("rotate", "0"))
        h2db = self._create_element("data-bind")
        h2db.set("_ref", page_data_bind_id)
        header2.append(h2db)
        h2bd = self._create_element("border")
        h2bd.set("_ref", border_id)
        header2.append(h2bd)
        h2bg = self._create_element("background")
        h2bg.set("_ref", page_bg_id)
        header2.append(h2bg)
        header2.append(self._create_element("height", "100"))
        header2.append(self._create_element("shrink", "false"))
        header2.append(self._create_element("children"))
        page2.append(header2)

        body2 = self._create_element("body")
        body2.append(self._create_element("id", "Body"))
        body2.append(self._create_element("parent-id", "page"))
        body2.append(self._create_element("z-index", "100"))
        body2.append(self._create_element("visible", "true"))
        body2.append(self._create_element("rotate", "0"))
        b2db = self._create_element("data-bind")
        b2db.set("_ref", page_data_bind_id)
        body2.append(b2db)
        b2bd = self._create_element("border")
        b2bd.set("_ref", border_id)
        body2.append(b2bd)
        b2bg = self._create_element("background")
        b2bg.set("_ref", page_bg_id)
        body2.append(b2bg)
        body2.append(self._create_element("shrink", "false"))

        body2_children = self._create_element("children")
        script_text = self._create_element("text")
        script_text.append(self._create_element("id", "text"))
        script_text.append(self._create_element("parent-id", "Body"))
        script_text.append(self._create_element("z-index", "100"))
        script_text.append(self._create_element("visible", "true"))
        script_text.append(self._create_element("rotate", "0"))
        script_text.append(self._create_element("adjust-height", "true"))
        script_text.append(self._create_element("appear-in-every-page", "false"))
        script_text.append(self._create_element("cross-page", "false"))
        sb = self._create_element("bound")
        sb.append(self._create_element("x", "0"))
        sb.append(self._create_element("y", "0"))
        sb.append(self._create_element("width", "2652"))
        sb.append(self._create_element("height", "1454"))
        script_text.append(sb)
        script_db = self._create_element("data-bind")
        script_db.set("_id", self._generate_id("DB"))
        script_bind = self._create_element("bind")
        script_bind.text = "//javascript\nJSON.stringify({})"
        script_db.append(script_bind)
        script_db.append(self._create_element("prepare-script"))
        script_db.append(self._create_element("post-script"))
        script_db.append(self._create_element("final-script"))
        script_text.append(script_db)
        script_border = self._create_element("border")
        script_border.set("_ref", border_id)
        script_text.append(script_border)
        script_background = self._create_element("background")
        script_background.set("_ref", page_bg_id)
        script_text.append(script_background)
        script_padding = self._create_element("padding")
        script_padding.set("_ref", padding_id)
        script_text.append(script_padding)
        script_font = self._create_font()
        script_text.append(script_font)
        script_para = self._create_element("paragraph")
        script_para.set("_ref", para_id)
        script_text.append(script_para)
        script_format = self._create_element("format")
        script_format.set("_ref", format_id)
        script_text.append(script_format)
        script_text.append(self._create_element("text-type", "TEXT"))
        script_text.append(self._create_element("use-rich-text-editor", "false"))
        body2_children.append(script_text)
        body2.append(body2_children)
        page2.append(body2)

        footer2 = self._create_element("footer")
        footer2.append(self._create_element("id", "Footer"))
        footer2.append(self._create_element("parent-id", "page"))
        footer2.append(self._create_element("z-index", "100"))
        footer2.append(self._create_element("visible", "true"))
        footer2.append(self._create_element("rotate", "0"))
        f2db = self._create_element("data-bind")
        f2db.set("_ref", page_data_bind_id)
        footer2.append(f2db)
        f2bd = self._create_element("border")
        f2bd.set("_ref", border_id)
        footer2.append(f2bd)
        f2bg = self._create_element("background")
        f2bg.set("_ref", page_bg_id)
        footer2.append(f2bg)
        footer2.append(self._create_element("height", "100"))
        footer2.append(self._create_element("shrink", "false"))
        footer2.append(self._create_element("children"))
        page2.append(footer2)

        pages.append(page2)
        template.append(pages)
        
        output_settings = self._create_element("output-settings")
        word_settings = self._create_element("word-settings")
        word_settings.append(self._create_element("body-reserved-height", "80"))
        output_settings.append(word_settings)
        pdf_settings = self._create_element("pdf-settings")
        pdf_settings.append(self._create_element("x-offset", "0"))
        pdf_settings.append(self._create_element("y-offset", "0"))
        output_settings.append(pdf_settings)
        template.append(output_settings)
        
        template.append(self._create_element("preview-data"))
        
        data_source = self._create_element("data-source-template")
        data_source.append(self._create_element("template-id"))
        template.append(data_source)
        
        return self._prettify(template)

    def _prettify(self, elem: ET.Element) -> str:
        rough_string = ET.tostring(elem, encoding='unicode')
        reparsed = minidom.parseString(rough_string)
        pretty = reparsed.toprettyxml(indent="\t", encoding=None)
        
        lines = pretty.split('\n')
        result_lines = []
        for line in lines:
            if line.strip():
                result_lines.append(line)
        
        result = '\n'.join(result_lines)
        
        import re
        result = re.sub(r'<([a-zA-Z0-9_-]+)\s*/>', r'<\1></\1>', result)
        result = re.sub(r'<([a-zA-Z0-9_-]+)((?:\s+[a-zA-Z0-9_-]+(?:="[^"]*")?)*)\s*/>', r'<\1\2></\1>', result)
        
        # 替换XML声明为固定格式
        result = re.sub(r'<\?xml[^?]*\?>', "<?xml version='1.0' encoding='utf-8'?>", result)
        
        return result


def convert_docx_to_xmreport(docx_path: str, output_path: str = None) -> str:
    """
    将docx文件转换为xmreport格式
    
    Args:
        docx_path: docx文件路径
        output_path: 输出文件路径（可选，默认为同目录下同名.xmreport文件）
    
    Returns:
        生成的xmreport文件路径
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"文件不存在: {docx_path}")
    
    ext = os.path.splitext(docx_path)[1].lower()
    
    if ext == '.doc':
        print(f"检测到.doc格式，正在转换为.docx...")
        docx_path = convert_doc_to_docx(docx_path)
        if output_path is None:
            output_path = os.path.splitext(docx_path)[0] + '.xmreport'
    elif ext != '.docx':
        raise ValueError(f"不支持的文件格式: {ext}，仅支持.doc和.docx格式")
    
    if output_path is None:
        output_path = os.path.splitext(docx_path)[0] + '.xmreport'
    
    print(f"正在解析: {docx_path}")
    parser = DocxParser(docx_path)
    page_content = parser.parse()
    page_settings = parser.get_page_settings()
    
    print(f"正在生成xmreport...")
    generator = XmreportGenerator()
    xml_content = generator.generate(page_content, page_settings)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(xml_content)
    
    print(f"转换完成: {output_path}")
    return output_path


def generate_empty_template(output_path: str = None) -> str:
    """
    生成一个空的骨架模板xmreport
    
    Args:
        output_path: 输出文件路径（可选，默认为当前目录下的empty_template.xmreport）
    
    Returns:
        生成的xmreport文件路径
    """
    if output_path is None:
        output_path = "empty_template.xmreport"
    
    print("正在生成空的骨架模板...")
    
    # 创建空的页面内容
    page_content = PageContent()
    page_content.header = []
    page_content.body = []
    page_content.footer = []
    
    # 创建默认页面设置（使用字典格式）
    page_settings = {
        'width': 1905,
        'height': 2777,
        'top_margin': 20,
        'bottom_margin': 118,
        'left_margin': 20,
        'right_margin': 20,
        'orientation': 'PORTRAIT'
    }
    
    # 生成xmreport
    generator = XmreportGenerator()
    xml_content = generator.generate(page_content, page_settings)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(xml_content)
    
    print(f"空模板生成完成: {output_path}")
    return output_path


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("用法: python docx_to_xmreport.py <doc/docx文件路径> [输出文件路径]")
        print("示例: python docx_to_xmreport.py report.docx")
        print("      python docx_to_xmreport.py report.doc output.xmreport")
        print("\n注意: .doc文件需要安装LibreOffice或Microsoft Word + pywin32")
        print("\n生成空的骨架模板: python docx_to_xmreport.py --empty [输出文件路径]")
        
        # 生成空模板
        output_file = sys.argv[2] if len(sys.argv) > 2 else None
        try:
            result = generate_empty_template(output_file)
            print(f"\n成功生成空模板: {result}")
        except Exception as e:
            print(f"生成空模板失败: {e}")
        sys.exit(1)
    
    # 检查是否生成空模板
    if sys.argv[1] == "--empty":
        output_file = sys.argv[2] if len(sys.argv) > 2 else None
        try:
            result = generate_empty_template(output_file)
            print(f"\n成功生成空模板: {result}")
        except Exception as e:
            print(f"生成空模板失败: {e}")
        sys.exit(0)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    try:
        result = convert_docx_to_xmreport(input_file, output_file)
        print(f"\n成功生成: {result}")
    except Exception as e:
        print(f"转换失败: {e}")
        sys.exit(1)
