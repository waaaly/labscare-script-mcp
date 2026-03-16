#!/usr/bin/env python3
"""
提取 Word 文档中的批注信息
用法: get-docx-comment <docx文件路径>
"""

import sys
from docx import Document
from docx.oxml.ns import qn


def extract_comments(file_path: str) -> dict:
    """提取 Word 文档中的所有批注信息"""
    doc = Document(file_path)

    # ── 1. 提取所有批注（id → 批注信息）──────────────────────────
    comments = {}
    try:
        comments_part = doc.part.part_related_by(
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
        )
        for comment in comments_part._element.findall(qn('w:comment')):
            cid = comment.get(qn('w:id'))
            author = comment.get(qn('w:author'))
            date = comment.get(qn('w:date'))
            text = ''.join(t.text for t in comment.findall('.//' + qn('w:t')) if t.text)
            comments[cid] = {'author': author, 'date': date, 'comment': text, 'ref_text': ''}
    except Exception:
        return None

    # ── 2. 收集所有段落（普通段落 + 表格单元格段落）────────────────
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)

    # ── 3. 遍历段落，提取批注标记范围内的被注释文本 ─────────────────
    for para in all_paragraphs:
        xml = para._element
        current_ids = set()
        for child in xml.iter():
            tag = child.tag
            if tag == qn('w:commentRangeStart'):
                current_ids.add(child.get(qn('w:id')))
            elif tag == qn('w:commentRangeEnd'):
                current_ids.discard(child.get(qn('w:id')))
            elif tag == qn('w:t') and child.text:
                for cid in current_ids:
                    if cid in comments:
                        comments[cid]['ref_text'] += child.text

    return comments


def print_comments(comments: dict) -> None:
    """打印批注信息"""
    if not comments:
        print("未找到任何批注")
        return

    print(f"共找到 {len(comments)} 条批注\n")
    print("=" * 60)
    for cid, info in comments.items():
        print(f"批注 ID   : {cid}")
        print(f"被注释文本 : {info['ref_text']}")
        print(f"批注内容   : {info['comment']}")
        print("-" * 60)


def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("用法: get-docx-comment <docx文件路径>")
        print("示例: get-docx-comment 'D:\\path\\to\\file.docx'")
        sys.exit(1)

    file_path = sys.argv[1]

    try:
        comments = extract_comments(file_path)
        if comments is None:
            print("该文档没有批注内容")
            sys.exit(0)
        print_comments(comments)
    except FileNotFoundError:
        print(f"错误: 文件不存在 - {file_path}")
        sys.exit(1)
    except Exception as e:
        print(f"错误: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
